from docx import Document
from bs4 import BeautifulSoup

def html_to_docx(html_string, output_path):
    # Tạo một tài liệu Word mới
    doc = Document()

    # Phân tích HTML và chèn nội dung vào tài liệu Word
    soup = BeautifulSoup(html_string, 'html.parser')
    for element in soup.descendants:
        if element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        # Thêm các điều kiện khác tùy thuộc vào cú pháp HTML của bạn

    # Lưu tài liệu Word
    doc.save(output_path)

# HTML input string
html_input = """
<div>
    <h1>Title</h1>
    <p>This is a paragraph.</p>
    <p>This is another paragraph.</p>
    <h2>Subtitle</h2>
    <p>This is a sub-paragraph.</p>
</div>
"""

# Đường dẫn tập tin Word output
output_path = "output.docx"

# Chuyển đổi HTML thành tập tin Word
html_to_docx(html_input, output_path)
