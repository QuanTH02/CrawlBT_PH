import time
from bs4 import BeautifulSoup
import docx
import pyautogui
import pyperclip
from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO

def html_to_docx(html_string, output_path):
    # Tạo một tài liệu Word mới
    doc = Document()

    # Phân tích HTML và chèn nội dung vào tài liệu Word
    soup = BeautifulSoup(html_string, 'html.parser')
    for element in soup.descendants:
        if element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'h4':
            doc.add_heading(element.text, level=4)
        elif element.name == 'h5':
            doc.add_heading(element.text, level=5)
        elif element.name == 'h6':
            doc.add_heading(element.text, level=6)
        elif element.name == 'ul':
            # Xử lý danh sách không sắp xếp
            ul = doc.add_paragraph()
            for li in element.find_all('li'):
                ul.add_run(f'{li.text}\n')
        elif element.name == 'ol':
            # Xử lý danh sách có sắp xếp
            ol = doc.add_paragraph()
            for li in element.find_all('li'):
                ol.add_run(f'{li.text}\n')
        elif element.name == 'a':
            # Xử lý liên kết
            link = doc.add_paragraph()
            link.add_run(element.text).bold = True
            link.add_run(f' ({element["href"]})')
        elif element.name == 'img':
            # Xử lý hình ảnh
            if 'src' in element.attrs:
                image_url = element['src']
                image_response = requests.get(image_url)
                if image_response.status_code == 200:
                    image_bytes = BytesIO(image_response.content)
                    image_run = doc.add_paragraph().add_run()
                    image_run.add_picture(image_bytes, width=Inches(1.0))
        elif element.name == 'table':
            # Xử lý bảng
            table = doc.add_table(rows=1, cols=1)
            for tr in element.find_all('tr'):
                row = table.add_row().cells
                for i, td in enumerate(tr.find_all(['td', 'th'])):
                    row[i].text = td.text
        elif element.name == 'code':
            # Xử lý mã nguồn
            code = doc.add_paragraph()
            code.add_run(element.text).font.family = 'Courier New'
        elif element.name == 'div':
            # Xử lý phần tử div
            div = doc.add_paragraph()
            div.add_run(element.text)
        elif element.name == 'span':
            # Xử lý phần tử span
            span = doc.add_paragraph()
            span.add_run(element.text)
        # Thêm các điều kiện khác tùy thuộc vào cú pháp HTML của bạn

    # Lưu tài liệu Word
    doc.save(output_path)

# Khởi tạo trình duyệt
driver = webdriver.Chrome()
driver.maximize_window()
page_main_url = "https://stackoverflow.com/questions/38899566/python-pyautogui-and-ctrl-c"
driver.get(page_main_url)
# time.sleep(1)

# pyautogui.keyDown('ctrl')
# pyautogui.keyDown('a')
# pyautogui.keyUp('a')
# pyautogui.keyUp('ctrl')

# time.sleep(1)

# pyautogui.keyDown('ctrl')
# pyautogui.keyDown('c')
# pyautogui.keyUp('c')
# pyautogui.keyUp('ctrl')

time.sleep(1)

div_element = driver.find_element(By.XPATH, "//*[@id=\"answer-68023685\"]")
script = "return arguments[0].innerHTML;"
content = driver.execute_script(script, div_element)

# Tạo tệp word và ghi nội dung đã sao chép vào đó
output_path = "output1.docx"

# Chuyển đổi HTML thành tập tin Word
html_to_docx(content, output_path)
# Đóng trình duyệt

time.sleep(1)
driver.quit()